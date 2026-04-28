from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/macbookpro/Developer/vachavox-macos/QA/app-screens-v311/vachavox-0.3.1.1-current-appearance-2026-04-27")
SCREENSHOTS = ROOT / "screenshots"
DOC_DIR = ROOT / "document"
DIAGRAMS = DOC_DIR / "diagrams"
MD_PATH = ROOT / "VachaVox-UI-visual-inventory-v311.md"
DOCX_PATH = DOC_DIR / "VachaVox-UI-visual-inventory-v311.docx"
MANIFEST_PATH = ROOT / "manifest.json"


screens = [
    {
        "file": "00_app-icon-resource.png",
        "surface": "App icon resource",
        "state": "Resource reference",
        "step": "Converted from Sources/VachaVox/Resources/VachaVox.icns.",
        "code": ["Sources/VachaVox/Resources/VachaVox.icns", "build/VachaVox.app/Contents/Info.plist"],
        "notes": "Static app icon resource converted to PNG with sips.",
    },
    {
        "file": "01_menu-bar-icon-idle.png",
        "surface": "Menu bar",
        "state": "Idle status item",
        "step": "Captured from the macOS menu bar while VachaVox was running.",
        "code": ["Sources/VachaVox/MenuBar/StatusItemController.swift", "Sources/VachaVox/Resources/MenuBarIcon.png"],
        "notes": "Menu bar crop includes adjacent menu extras because status-item direct click/crop access was limited.",
    },
    {
        "file": "10_settings-home.png",
        "surface": "Settings",
        "state": "Home tab, model ready",
        "step": "Open VachaVox Settings, select Home.",
        "code": ["Sources/VachaVox/UI/SettingsView.swift", "Sources/VachaVox/App/DictationCoordinator.swift"],
        "notes": "Shows model readiness, primary dictation action, load/refresh controls, and summary rows.",
    },
    {
        "file": "11_settings-models.png",
        "surface": "Settings",
        "state": "Models tab",
        "step": "Open VachaVox Settings, select Models.",
        "code": ["Sources/VachaVox/UI/SettingsView.swift", "Sources/VachaVox/Models/ModelCatalog.swift", "Sources/VachaVox/Models/ModelDownloadService.swift"],
        "notes": "Shows selected model summary, two installed models, and available downloads.",
    },
    {
        "file": "12_settings-dictation.png",
        "surface": "Settings",
        "state": "Dictation tab",
        "step": "Open VachaVox Settings, select Dictation.",
        "code": ["Sources/VachaVox/UI/SettingsView.swift", "Sources/VachaVox/Settings/AppSettings.swift"],
        "notes": "Shows output segmented control, punctuation toggle, silence slider, and performance picker.",
    },
    {
        "file": "13_settings-hotkeys.png",
        "surface": "Settings",
        "state": "Hotkeys tab",
        "step": "Open VachaVox Settings, select Hotkeys.",
        "code": ["Sources/VachaVox/UI/SettingsView.swift", "Sources/VachaVox/MenuBar/HotKeyService.swift", "Sources/VachaVox/Settings/AppSettings.swift"],
        "notes": "Shows push-to-talk mode and Command-Shift-D preset.",
    },
    {
        "file": "14_settings-permissions.png",
        "surface": "Settings",
        "state": "Permissions tab",
        "step": "Open VachaVox Settings, select Permissions.",
        "code": ["Sources/VachaVox/UI/SettingsView.swift", "Sources/VachaVox/Permissions/PermissionsService.swift", "Sources/VachaVox/Utilities/SystemSettingsOpener.swift"],
        "notes": "Shows microphone granted, accessibility trusted, privacy settings buttons, and start-at-login toggle.",
    },
    {
        "file": "15_settings-privacy-about.png",
        "surface": "Settings",
        "state": "Privacy/About tab",
        "step": "Open VachaVox Settings, select Privacy/About.",
        "code": ["Sources/VachaVox/UI/SettingsView.swift"],
        "notes": "Shows local privacy promise and dependency/about notes.",
    },
    {
        "file": "16_settings-model-delete-confirmation.png",
        "surface": "Settings",
        "state": "Model delete confirmation",
        "step": "Models tab, click Delete for Whisper Small, then cancel.",
        "code": ["Sources/VachaVox/UI/SettingsView.swift", "Sources/VachaVox/App/DictationCoordinator.swift", "Sources/VachaVox/Models/ModelCatalog.swift"],
        "notes": "Confirmation only; no model folders were deleted.",
    },
    {
        "file": "17_settings-hotkeys-custom-recorder.png",
        "surface": "Settings",
        "state": "Custom shortcut recorder",
        "step": "Hotkeys tab, change Preset to Custom shortcut; setting restored afterward.",
        "code": ["Sources/VachaVox/UI/SettingsView.swift", "Sources/VachaVox/MenuBar/HotKeyService.swift", "Sources/VachaVox/Settings/AppSettings.swift"],
        "notes": "Shows KeyboardShortcuts recorder field. The preset was restored to Command-Shift-D after capture.",
    },
]

skipped = [
    {
        "file": "02_menu-bar-popover-idle.png",
        "reason": "The LSUIElement status item was visible in the menu bar, but direct status-item activation was blocked from this automation environment after the Settings window was closed.",
    },
    {
        "file": "03_menu-bar-popover-paused.png",
        "reason": "Depends on accessing the menu bar popover toggle; skipped for the same status-item activation limitation.",
    },
    {
        "file": "04_recording-overlay-listening.png",
        "reason": "The overlay is reachable through Start Dictation, but the only exposed entry point after capture was the inaccessible status item; no permission state was changed to force it.",
    },
    {
        "file": "05_recording-overlay-transcribing.png",
        "reason": "Skipped because it depends on a live recording stop/transcription cycle; no synthetic audio or model operation was forced.",
    },
    {
        "file": "06_preview-dictation-window.png",
        "reason": "Skipped because Preview output requires completing a dictation flow; no text output state was forced.",
    },
    {
        "file": "18_settings-permissions-system-settings-link.png",
        "reason": "Skipped to avoid opening System Settings or changing privacy permission state.",
    },
    {
        "file": "19_popover-error-or-missing-model.png",
        "reason": "Skipped because the current environment had a ready installed model; no model folders were moved or removed to create an error.",
    },
    {
        "file": "20_popover-last-transcript.png",
        "reason": "Skipped because it requires a completed dictation session and popover access.",
    },
]

flows = {
    "App Launch And Menu Bar Lifecycle": """flowchart TD
  Launch["Launch VachaVox"] --> AppDelegate["AppDelegate creates services"]
  AppDelegate --> Model["AppModel loads saved settings"]
  Model --> StatusItem["StatusItemController installs menu bar item"]
  StatusItem --> Popover["PopoverView opens from status item"]
  Popover --> Settings["Open Settings window"]
  AppDelegate --> LoadModel["Load best available local model"]
  LoadModel --> Ready["Status: Model ready"]
""",
    "First-Run Permissions And Model Readiness": """flowchart TD
  Start["Start Dictation"] --> Mic["Request microphone access"]
  Mic -->|Granted| Capture["Begin audio capture"]
  Mic -->|Denied| Error["Show microphone required error"]
  Settings["Permissions tab"] --> SystemSettings["Open macOS privacy panes"]
  Models["Models tab"] --> Download["Download or reveal local model"]
  Download --> Ready["Installed or ready model status"]
""",
    "Dictation Flow": """flowchart TD
  Trigger["Hotkey or Start Dictation"] --> Listening["Phase: listening"]
  Listening --> Stop["Stop Dictation"]
  Stop --> Trim["VoiceActivityService trims silence"]
  Trim --> Prepare["Prepare selected transcription engine"]
  Prepare --> Transcribe["Transcribe audio"]
  Transcribe --> Deliver["TextOutputService delivers text"]
  Deliver --> Idle["Return to idle with status"]
""",
    "Settings Navigation Flow": """flowchart TD
  Settings["VachaVox Settings"] --> Home["Home"]
  Settings --> Models["Models"]
  Settings --> Dictation["Dictation"]
  Settings --> Hotkeys["Hotkeys"]
  Settings --> Permissions["Permissions"]
  Settings --> Privacy["Privacy/About"]
""",
    "Model Management Flow": """flowchart TD
  Catalog["ModelCatalog descriptors"] --> Scan["ModelStore scans local folders"]
  Scan --> Installed["Downloaded Models"]
  Scan --> Missing["Available Downloads"]
  Installed --> Load["Load selected model"]
  Installed --> DeleteConfirm["Delete confirmation"]
  Missing --> Download["Download model"]
""",
    "Output Delivery Flow": """flowchart TD
  Transcript["Transcript text"] --> Mode{"Output mode"}
  Mode --> Copy["Copy to clipboard"]
  Mode --> Paste["Paste at focused target"]
  Mode --> Preview["Show Preview Dictation window"]
  Preview --> Accept["Copy edited text"]
""",
}


def ensure_dirs() -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    DIAGRAMS.mkdir(parents=True, exist_ok=True)


def draw_flow(title: str, lines: list[str], path: Path) -> None:
    width = 1400
    height = 260 + 86 * len(lines)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 36)
        body_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 25)
    except OSError:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
    draw.rectangle([0, 0, width, height], fill=(250, 252, 255))
    draw.text((54, 36), title, fill=(24, 49, 83), font=title_font)
    y = 118
    x = 58
    for i, line in enumerate(lines):
        box = [x, y, width - 90, y + 52]
        draw.rounded_rectangle(box, radius=16, fill=(238, 244, 252), outline=(169, 190, 220), width=2)
        draw.text((x + 24, y + 12), line, fill=(31, 41, 55), font=body_font)
        if i < len(lines) - 1:
            cx = width // 2
            draw.line([cx, y + 54, cx, y + 82], fill=(74, 105, 150), width=4)
            draw.polygon([(cx - 8, y + 76), (cx + 8, y + 76), (cx, y + 90)], fill=(74, 105, 150))
        y += 86
    image.save(path)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.replace("/", "/ ").split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = word if not current else f"{current} {word}"
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def draw_table_image(title: str, headers: list[str], rows: list[list[str]], path: Path, col_widths: list[int]) -> None:
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 30)
        header_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 20)
        body_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 19)
    except OSError:
        title_font = header_font = body_font = ImageFont.load_default()

    pad_x = 18
    pad_y = 14
    line_h = 25
    width = sum(col_widths) + pad_x * 2
    probe = Image.new("RGB", (width, 100), "white")
    draw = ImageDraw.Draw(probe)
    row_heights: list[int] = []
    wrapped: list[list[list[str]]] = []
    for row in [headers] + rows:
        row_wrap: list[list[str]] = []
        max_lines = 1
        for i, cell in enumerate(row):
            font = header_font if row == headers else body_font
            lines = wrap_text(draw, cell, font, col_widths[i] - 18)
            row_wrap.append(lines)
            max_lines = max(max_lines, len(lines))
        wrapped.append(row_wrap)
        row_heights.append(max(44, pad_y * 2 + max_lines * line_h))

    height = 74 + sum(row_heights) + pad_y * 2
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle([0, 0, width, height], fill=(250, 252, 255))
    draw.text((pad_x, 20), title, fill=(24, 49, 83), font=title_font)
    y = 70
    for r_idx, row in enumerate(wrapped):
        fill = (226, 235, 247) if r_idx == 0 else ((255, 255, 255) if r_idx % 2 else (245, 248, 252))
        draw.rectangle([pad_x, y, width - pad_x, y + row_heights[r_idx]], fill=fill, outline=(180, 193, 210))
        x = pad_x
        for c_idx, lines in enumerate(row):
            draw.line([x, y, x, y + row_heights[r_idx]], fill=(180, 193, 210), width=1)
            font = header_font if r_idx == 0 else body_font
            text_y = y + 11
            for line in lines:
                draw.text((x + 9, text_y), line, fill=(31, 41, 55), font=font)
                text_y += line_h
            x += col_widths[c_idx]
        draw.line([width - pad_x, y, width - pad_x, y + row_heights[r_idx]], fill=(180, 193, 210), width=1)
        y += row_heights[r_idx]
    image.save(path)


def make_diagrams() -> dict[str, Path]:
    diagram_steps = {
        "App Launch And Menu Bar Lifecycle": [
            "Launch VachaVox",
            "AppDelegate creates app model, services, status item, overlay",
            "Status item opens PopoverView",
            "Settings opens from popover gear",
            "Best available model loads and reports readiness",
        ],
        "First-Run Permissions And Model Readiness": [
            "User starts dictation or opens Permissions",
            "Microphone and accessibility state are checked",
            "Models tab scans local model folders",
            "Missing models can be downloaded, installed models can load",
        ],
        "Dictation Flow": [
            "Hotkey or Start Dictation",
            "Listening phase with recording overlay",
            "Stop, trim silence, prepare engine",
            "Transcribe locally, deliver text, return to idle",
        ],
        "Settings Navigation Flow": [
            "Settings window",
            "Home, Models, Dictation, Hotkeys, Permissions, Privacy/About",
        ],
        "Model Management Flow": [
            "ModelCatalog defines known models",
            "ModelStore validates local folders",
            "Installed models can select/load/reveal/delete",
            "Missing downloadable models expose Download and Open Source",
        ],
        "Output Delivery Flow": [
            "Transcript text from transcription engine",
            "Copy mode writes clipboard",
            "Paste mode inserts at focused target",
            "Preview mode opens editable Preview Dictation window",
        ],
    }
    out = {}
    for idx, (title, lines) in enumerate(diagram_steps.items(), 1):
        path = DIAGRAMS / f"flow-{idx:02d}-{title.lower().replace(' ', '-')}.png"
        draw_flow(title, lines, path)
        out[title] = path
    return out


def make_table_images() -> dict[str, Path]:
    visual_path = DIAGRAMS / "table-visual-inventory.png"
    draw_table_image(
        "Visual Inventory",
        ["Screenshot", "Surface", "State", "Notes"],
        [[s["file"], s["surface"], s["state"], s["notes"]] for s in screens],
        visual_path,
        [260, 170, 270, 660],
    )
    skipped_path = DIAGRAMS / "table-skipped-captures.png"
    draw_table_image(
        "Skipped Captures",
        ["Planned screenshot", "Reason"],
        [[i["file"], i["reason"]] for i in skipped],
        skipped_path,
        [360, 980],
    )
    code_path = DIAGRAMS / "table-code-mapping.png"
    draw_table_image(
        "Code Mapping Notes",
        ["File", "UI responsibility"],
        [
            ["SettingsView.swift", "Sidebar tabs, settings panels, model cards, status badges, permission controls, and SettingsWindowController."],
            ["PopoverView.swift", "Menu bar popover header, level meter, primary action, readiness rows, transcript card, Paused toggle, Settings, and Quit."],
            ["RecordingOverlayWindowController.swift", "Floating recording/transcribing panel shown during transient dictation states."],
            ["PreviewWindowController.swift", "Editable transcript preview window for Preview output mode."],
            ["AppSettings.swift", "Persisted settings options and defaults."],
            ["DictationCoordinator.swift", "Runtime state transitions, model operations, permission checks, transcription, and output delivery."],
            ["ModelCatalog.swift", "Model metadata, local paths, model statuses, validation, and best-available priority."],
        ],
        code_path,
        [360, 980],
    )
    return {"visual": visual_path, "skipped": skipped_path, "code": code_path}


def write_manifest() -> None:
    data = {
        "app": "VachaVox",
        "version": "0.3.1.1",
        "capture_date": "2026-04-27",
        "source_repo": "/Users/macbookpro/Developer/vachavox-macos",
        "app_bundle": "/Users/macbookpro/Developer/vachavox-macos/build/VachaVox.app",
        "screenshots": screens,
        "skipped_screenshots": skipped,
    }
    MANIFEST_PATH.write_text(json.dumps(data, indent=2), encoding="utf-8")


def md_image(name: str) -> str:
    return f"![{name}]({(SCREENSHOTS / name).as_posix()})"


def write_markdown() -> None:
    lines: list[str] = []
    lines.append("# VachaVox UI Visual Inventory v0.3.1.1")
    lines.append("")
    lines.append("Capture date: 2026-04-27")
    lines.append("")
    lines.append("VachaVox is a local-first macOS menu bar dictation app. It records only while dictation is active, transcribes with local Parakeet or WhisperKit models, and delivers text through Copy, Paste, or Preview output modes. The app surface is intentionally compact: the menu bar popover is the fast path, while Settings handles model management, dictation behavior, hotkeys, permissions, and privacy/about information.")
    lines.append("")
    lines.append("## Visual Inventory")
    lines.append("")
    lines.append("| Screenshot | Surface | State | Related code | Capture notes |")
    lines.append("|---|---|---|---|---|")
    for s in screens:
        code = "<br>".join(f"`{c}`" for c in s["code"])
        lines.append(f"| `{s['file']}` | {s['surface']} | {s['state']} | {code} | {s['notes']} |")
    lines.append("")
    lines.append("## Skipped Screenshots")
    lines.append("")
    for item in skipped:
        lines.append(f"- `{item['file']}`: {item['reason']}")
    lines.append("")
    lines.append("## User Flow Overview")
    lines.append("")
    for title, flow in flows.items():
        lines.append(f"### {title}")
        lines.append("")
        lines.append("```mermaid")
        lines.append(flow.strip())
        lines.append("```")
        lines.append("")
    lines.append("## Per-Screen Notes")
    lines.append("")
    per_screen = {
        "00_app-icon-resource.png": ("App Icon Resource", "Identifies the product in app bundle resources and the Settings/Home header icon.", "Bundle resource and Info.plist icon lookup.", "Large square app icon with VachaVox branding.", "Keep a consistent icon treatment between app icon, menu bar glyph, and in-app brand mark."),
        "01_menu-bar-icon-idle.png": ("Menu Bar Icon - Idle", "Provides the always-available entry point for the app.", "Run app; inspect macOS menu bar.", "Template status icon supplied by MenuBarIcon.png while idle.", "Status item needs a clearer active/paused/error visual language if more states are added."),
        "10_settings-home.png": ("Settings - Home", "Gives an operational summary and direct model/dictation actions.", "Open Settings, select Home.", "Start Dictation, Load Model, Refresh Models, selected model, engine, output, hotkey, status.", "Home works as a dashboard but could use stronger grouping for readiness, action, and configuration summary."),
        "11_settings-models.png": ("Settings - Models", "Centralizes local model status and model operations.", "Open Settings, select Models.", "Selected model summary, downloaded models, available downloads, status badges, select/load/delete/reveal/source actions.", "Model cards are functional but dense; actions could be grouped by primary, file management, and external source."),
        "12_settings-dictation.png": ("Settings - Dictation", "Controls output delivery and transcription behavior.", "Open Settings, select Dictation.", "Output segmented control, punctuation toggle, silence sensitivity slider, performance picker.", "Add short but compact labels for sensitivity/performance consequences if redesign has room."),
        "13_settings-hotkeys.png": ("Settings - Hotkeys", "Controls trigger behavior and shortcut preset.", "Open Settings, select Hotkeys.", "Mode segmented control, preset picker, implementation note.", "The low-level Fn note is useful but visually quiet; redesign could make shortcut mode and preset easier to scan."),
        "14_settings-permissions.png": ("Settings - Permissions", "Shows privacy permission readiness and links to macOS privacy panes.", "Open Settings, select Permissions.", "Microphone status, accessibility status, settings buttons, start-at-login toggle.", "Permission status could use badges and remediation language for denied/not trusted states."),
        "15_settings-privacy-about.png": ("Settings - Privacy/About", "States local processing and dependency/about information.", "Open Settings, select Privacy/About.", "Local privacy promise and dependency list.", "This tab is sparse; could combine privacy assurance, version, licenses, and diagnostics in a structured layout."),
        "16_settings-model-delete-confirmation.png": ("Settings - Delete Confirmation", "Protects local model folders from accidental deletion.", "Models tab, Delete on Whisper Small, then Cancel.", "Confirmation dialog with destructive action and target folder path.", "Long paths wrap heavily; redesign could separate model name, impact, and path in clearer rows."),
        "17_settings-hotkeys-custom-recorder.png": ("Settings - Custom Shortcut Recorder", "Shows the custom shortcut capture state.", "Hotkeys tab, choose Custom shortcut.", "KeyboardShortcuts recorder with existing shortcut and clear control.", "Recorder fits the panel, but current shortcut inheritance should be clearer during preset switching."),
    }
    for s in screens:
        title, purpose, entry, controls, notes = per_screen[s["file"]]
        lines.append(f"### {title}")
        lines.append("")
        lines.append(md_image(s["file"]))
        lines.append("")
        lines.append(f"Purpose: {purpose}")
        lines.append("")
        lines.append(f"Entry path/user steps: {entry}")
        lines.append("")
        lines.append(f"Visible controls/options: {controls}")
        lines.append("")
        lines.append(f"Current behavior: {s['notes']}")
        lines.append("")
        lines.append("Related code files/folders:")
        for c in s["code"]:
            lines.append(f"- `{c}`")
        lines.append("")
        lines.append(f"UI rework notes and design-system opportunities: {notes}")
        lines.append("")
    lines.append("## Code Mapping Notes")
    lines.append("")
    mapping = [
        ("SettingsView.swift", "Defines NavigationSplitView sidebar tabs, settings panels, model cards, status badges, confirmation dialog, permissions pane, and SettingsWindowController."),
        ("PopoverView.swift", "Defines compact menu bar surface with app header, level meter, primary Start/Stop button, readiness rows, transcript card, Paused toggle, Settings, and Quit actions."),
        ("RecordingOverlayWindowController.swift", "Defines transient floating NSPanel for listening/transcribing states and positions it near the focused caret or mouse."),
        ("PreviewWindowController.swift", "Defines editable Preview Dictation window with TextEditor, Cancel, and Copy accept action."),
        ("AppSettings.swift", "Stores selected engine/model, hotkey mode and preset, output mode, punctuation, silence sensitivity, performance mode, launch-at-login, paused state, and hotkey description."),
        ("DictationCoordinator.swift", "Owns dictation state transitions, permission checks, recording start/stop, local transcription, output delivery, model load/download/delete/reveal, and status updates."),
        ("ModelCatalog.swift", "Defines Parakeet and WhisperKit model metadata, validation, default model, priority order, local model paths, and status labels."),
    ]
    for file, note in mapping:
        lines.append(f"- `{file}`: {note}")
    lines.append("")
    lines.append("## UI Rework Checklist")
    lines.append("")
    checklist = [
        "Navigation hierarchy: decide whether Home is a dashboard, setup checklist, or status summary.",
        "Control consistency: align segmented controls, pop-up buttons, toggles, and destructive actions with a predictable settings grammar.",
        "Status/state language: normalize Ready, Installed, Missing, Loading, Downloading, Paused, and error messaging across popover and Settings.",
        "Empty/error/loading states: design explicit missing-model, permission-denied, loading-model, and failed-transcription states.",
        "Model management density: separate model identity, status, file path, and operations so dense cards remain scannable.",
        "Permission guidance: add concise remediation for microphone/accessibility when not granted.",
        "Accessibility and keyboard behavior: verify sidebar navigation, shortcut recorder focus, destructive confirmation, and VoiceOver labels.",
        "Visual design tokens: define spacing, typography, radius, color, icons, and status badges shared by popover, settings cards, and overlays.",
    ]
    for item in checklist:
        lines.append(f"- {item}")
    lines.append("")
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(24, 49, 83)


def set_cell_width(cell, width: float) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths is None:
        if len(headers) == 4:
            widths = [1.25, 1.0, 1.35, 3.0]
        elif len(headers) == 2:
            widths = [1.9, 4.7]
        else:
            widths = [6.6 / len(headers)] * len(headers)
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_width(table.rows[0].cells[i], widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_width(cells[i], widths[i])
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph("")


def add_image(doc: Document, path: Path, caption: str, width_inches: float = 6.3) -> None:
    doc.add_picture(str(path), width=Inches(width_inches))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def write_docx(diagrams: dict[str, Path], table_images: dict[str, Path]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VachaVox UI Visual Inventory")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(24, 49, 83)
    subtitle = doc.add_paragraph("Current appearance review for VachaVox 0.3.1.1 | Captured 2026-04-27")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    doc.add_paragraph(
        "This report captures the current VachaVox macOS UI for future redesign planning. "
        "The Markdown artifact is the canonical source; this Word version packages the same context with readable screenshots, tables, and rendered flow diagrams."
    )

    add_heading(doc, "Visual Inventory", 1)
    add_image(doc, table_images["visual"], "Visual inventory table", width_inches=6.6)

    add_heading(doc, "Flow Overview", 1)
    for title, path in diagrams.items():
        add_heading(doc, title, 2)
        add_image(doc, path, title, width_inches=6.4)

    add_heading(doc, "Per-Screen Inventory", 1)
    for s in screens:
        add_heading(doc, s["state"], 2)
        doc.add_paragraph(f"Surface: {s['surface']}")
        doc.add_paragraph(f"Entry path: {s['step']}")
        doc.add_paragraph(f"Related code: {', '.join(s['code'])}")
        add_image(doc, SCREENSHOTS / s["file"], s["file"], width_inches=6.4 if s["file"] != "00_app-icon-resource.png" else 2.2)
        doc.add_paragraph(s["notes"])

    add_heading(doc, "Skipped Captures", 1)
    add_image(doc, table_images["skipped"], "Skipped capture table", width_inches=6.6)

    add_heading(doc, "Code Mapping Notes", 1)
    add_image(doc, table_images["code"], "Code mapping table", width_inches=6.6)

    add_heading(doc, "UI Rework Checklist", 1)
    for item in [
        "Clarify Settings Home as dashboard, setup checklist, or status summary.",
        "Normalize status badges and language across Settings and popover.",
        "Reduce model card density while preserving advanced actions.",
        "Design explicit empty, error, loading, paused, permission-denied, and missing-model states.",
        "Audit keyboard navigation, VoiceOver labels, and shortcut recorder behavior.",
        "Define shared tokens for spacing, typography, radius, iconography, and status colors.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main() -> None:
    ensure_dirs()
    diagrams = make_diagrams()
    table_images = make_table_images()
    write_manifest()
    write_markdown()
    write_docx(diagrams, table_images)
    print(MD_PATH)
    print(DOCX_PATH)
    print(MANIFEST_PATH)


if __name__ == "__main__":
    main()
